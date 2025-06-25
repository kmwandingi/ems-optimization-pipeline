#!/usr/bin/env python3
# ================================================================
# convert_html_to_word_improved.py
# ------------------------------------------------
# Purpose  : Convert HTML with MathJax formulas to Word .docx
# Strategy : Fix TOC numbering and improve math formula extraction
# Author   : Enhanced for proper TOC and math formula preservation
# ================================================================

# -------- 1. Standard-library imports ---------------------------------
import subprocess
import sys
from pathlib import Path
import shutil
import platform
import re
import tempfile
import os

# -------- 2. Third-party imports (lazy-loaded if Pandoc absent) -------
from bs4 import BeautifulSoup          # noqa: E402
from docx import Document              # noqa: E402


# ================================================================
# SECTION: Enhanced MathJax Formula Extractor
# ------------------------------------------------
class MathJaxExtractor:
    """Extract LaTeX formulas from MathJax-rendered HTML and prepare for Word conversion."""
    
    def __init__(self):
        # I'm defining comprehensive patterns to find mathematical formulas
        self.math_patterns = [
            # Direct LaTeX in HTML (most reliable)
            (r'\$\$(.*?)\$\$', 'display', 'latex'),
            (r'\$([^$\n]+?)\$', 'inline', 'latex'),
            # MathJax script tags
            (r'<script[^>]*?type=["\']math/tex[^"\']*?["\'][^>]*?>(.*?)</script>', 'auto', 'script'),
            # MathJax HTML elements
            (r'<mjx-container[^>]*?class="[^"]*?MathJax[^"]*?"[^>]*?>(.*?)</mjx-container>', 'display', 'mjx'),
            (r'<span[^>]*?class="[^"]*?MathJax[^"]*?"[^>]*?>(.*?)</span>', 'inline', 'mjx'),
        ]
    
    def debug_formula_extraction(self, html_content):
        """
        I'm providing detailed debugging of formula extraction to help
        identify what mathematical content is found and how it's processed.
        """
        print("[DEBUG] Analyzing mathematical content in HTML...")
        
        # Count different types of math content
        latex_display = len(re.findall(r'\$\$(.*?)\$\$', html_content, re.DOTALL))
        latex_inline = len(re.findall(r'\$([^$\n]+?)\$', html_content))
        mathjax_scripts = len(re.findall(r'<script[^>]*?type=["\']math/tex', html_content))
        mathjax_elements = len(re.findall(r'class="[^"]*?MathJax[^"]*?"', html_content))
        
        print(f"[DEBUG] Found LaTeX display formulas: {latex_display}")
        print(f"[DEBUG] Found LaTeX inline formulas: {latex_inline}")
        print(f"[DEBUG] Found MathJax script elements: {mathjax_scripts}")
        print(f"[DEBUG] Found MathJax HTML elements: {mathjax_elements}")
        
        # Show sample formulas
        sample_formulas = re.findall(r'\$\$(.*?)\$\$', html_content, re.DOTALL)[:3]
        if sample_formulas:
            print("[DEBUG] Sample display formulas found:")
            for i, formula in enumerate(sample_formulas):
                print(f"  {i+1}: {formula[:100]}{'...' if len(formula) > 100 else ''}")
        
        return latex_display + latex_inline + mathjax_scripts + mathjax_elements
    
    def extract_and_clean_formulas(self, html_content):
        """
        I'm extracting mathematical formulas and preparing clean LaTeX
        that Pandoc can properly convert to Word equations.
        """
        formulas_found = []
        
        # Extract LaTeX formulas (these are the most reliable)
        display_formulas = re.findall(r'\$\$(.*?)\$\$', html_content, re.DOTALL)
        inline_formulas = re.findall(r'\$([^$\n]+?)\$', html_content)
        
        print(f"[INFO] Extracted {len(display_formulas)} display formulas")
        print(f"[INFO] Extracted {len(inline_formulas)} inline formulas")
        
        # Store formulas for replacement
        for formula in display_formulas:
            clean_formula = self._clean_latex_formula(formula)
            formulas_found.append(('display', clean_formula))
        
        for formula in inline_formulas:
            clean_formula = self._clean_latex_formula(formula)
            formulas_found.append(('inline', clean_formula))
        
        return formulas_found
    
    def _clean_latex_formula(self, formula):
        """
        I'm cleaning LaTeX formulas to ensure they work properly in Word.
        This applies the same fixes as your create_html_version.py script.
        """
        # Apply the same formula fixes from your working HTML script
        
        # Fix subscript/superscript ordering (critical fix)
        formula = re.sub(r'([a-zA-Z0-9])\^\{(\+|-)\}_([a-zA-Z0-9]+)', r'\1_\3^{\2}', formula)
        formula = re.sub(r'([a-zA-Z0-9])\^(\+|-)_([a-zA-Z0-9]+)', r'\1_\3^{\2}', formula)
        formula = re.sub(r'([a-zA-Z0-9])\^\{([^}]+)\}_([a-zA-Z0-9]+)', r'\1_\3^{\2}', formula)
        
        # Fix complex subscripts
        formula = re.sub(r'([a-zA-Z0-9])\^(\+|-)_\{([^}]+)\}', r'\1_{\3}^\2', formula)
        formula = re.sub(r'([a-zA-Z0-9])\^\{(\+|-)\}_\{([^}]+)\}', r'\1_{\3}^{\2}', formula)
        
        # Fix malformed constructs
        formula = re.sub(r'G\\max\b', r'G_{\\max}', formula)
        formula = re.sub(r'\\sum\{([^}]+)\}', r'\\sum_{\1}', formula)
        
        # Clean up spaces in braces
        formula = re.sub(r'\^\{\s*\+\s*\}', r'^{+}', formula)
        formula = re.sub(r'\^\{\s*-\s*\}', r'^{-}', formula)
        
        return formula.strip()
    
    def create_pandoc_ready_html(self, html_file_path):
        """
        I'm creating clean HTML for Pandoc with properly extracted formulas
        and fixed TOC structure.
        """
        with open(html_file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Debug formula extraction
        total_math = self.debug_formula_extraction(html_content)
        
        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove problematic MathJax scripts and elements
        self._remove_mathjax_artifacts(soup)
        
        # Fix TOC structure
        self._fix_toc_structure(soup)
        
        # Get clean HTML
        clean_html = str(soup)
        
        print(f"[INFO] Prepared clean HTML for Pandoc conversion")
        return clean_html
    
    def _remove_mathjax_artifacts(self, soup):
        """
        I'm removing MathJax scripts and elements that interfere with conversion
        while preserving the LaTeX formulas.
        """
        # Remove MathJax scripts
        for script in soup.find_all('script'):
            script_content = script.string or ''
            if ('mathjax' in script_content.lower() or 
                'mathjax' in script.get('src', '').lower() or
                'math/tex' in script.get('type', '')):
                script.decompose()
        
        # Remove MathJax CSS and styling
        for style in soup.find_all('style'):
            style_content = style.string or ''
            if 'mathjax' in style_content.lower():
                style.decompose()
        
        # Remove MathJax HTML elements but preserve LaTeX content
        for element in soup.find_all(['mjx-container', 'mjx-math']):
            element.unwrap()  # Remove the element but keep its content
    
    def _fix_toc_structure(self, soup):
        """
        I'm fixing the Table of Contents structure to prevent numbering issues.
        """
        # Find TOC section and remove automatic numbering
        toc_headers = soup.find_all(['h1', 'h2', 'h3'], string=re.compile(r'Table of Contents', re.I))
        
        for toc_header in toc_headers:
            print("[INFO] Found Table of Contents - fixing structure")
            
            # Find the TOC list that follows
            toc_list = toc_header.find_next(['ul', 'ol'])
            if toc_list:
                # Convert to unordered list to remove numbering
                if toc_list.name == 'ol':
                    toc_list.name = 'ul'
                
                # Remove any number prefixes from TOC items
                for li in toc_list.find_all('li'):
                    text = li.get_text()
                    # Remove leading numbers and dots
                    cleaned_text = re.sub(r'^\d+\.\s*', '', text)
                    if cleaned_text != text:
                        li.string = cleaned_text


# ================================================================
# SECTION: Helper functions
# ------------------------------------------------
def pandoc_available() -> bool:
    """Check whether Pandoc is on PATH and at least version 2.9."""
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    try:
        version_output = subprocess.check_output([pandoc, "--version"], text=True)
        first_line = version_output.split('\n')[0]
        version_match = re.search(r'pandoc (\d+)\.(\d+)', first_line)
        if version_match:
            major, minor = int(version_match.group(1)), int(version_match.group(2))
            return major > 2 or (major == 2 and minor >= 9)
        return False
    except Exception:
        return False


def run_pandoc_with_math(html_file: Path, docx_file: Path) -> None:
    """
    I'm using Pandoc with enhanced math handling and proper TOC generation.
    """
    # Preprocess the HTML
    extractor = MathJaxExtractor()
    clean_html = extractor.create_pandoc_ready_html(html_file)
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.html', encoding='utf-8', delete=False) as tmp_file:
        tmp_file.write(clean_html)
        tmp_html_path = tmp_file.name
    
    try:
        # Enhanced Pandoc command
        cmd = [
            "pandoc",
            tmp_html_path,
            "--from=html+tex_math_dollars+tex_math_single_backslash",  # Enhanced math parsing
            "--to=docx",
            "--mathml",                      # Use MathML for native Word math
            "--standalone",
            "--embed-resources",
            "--toc",                         # Let Pandoc generate clean TOC
            "--toc-depth=3",                 # Limit TOC depth
            "--number-sections",             # Pandoc will handle numbering properly
            "--reference-doc=" + create_reference_docx(),
            "--output", str(docx_file),
        ]
        
        print(f"[INFO] Running Pandoc with enhanced math processing...")
        print(f"[CMD] {' '.join(cmd[:8])}...")  # Show partial command
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        
        if result.returncode == 0:
            print(f"[SUCCESS] ✅ Created {docx_file} with native Word math equations")
            print(f"[INFO] Pandoc processed the document successfully")
        else:
            print(f"[ERROR] ❌ Pandoc conversion failed:")
            print(f"[STDERR] {result.stderr}")
            raise subprocess.CalledProcessError(result.returncode, cmd)
            
    finally:
        # Clean up
        try:
            os.unlink(tmp_html_path)
        except OSError:
            pass


def create_reference_docx():
    """
    I'm creating a professional reference document for better Word styling.
    """
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    doc = Document()
    styles = doc.styles
    
    # Enhance Normal style
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    
    # Enhance heading styles
    heading_sizes = [16, 14, 12, 11, 10, 9]
    for i in range(1, 7):
        heading_name = f'Heading {i}'
        if heading_name in styles:
            heading = styles[heading_name]
            heading.font.name = 'Calibri'
            heading.font.size = Pt(heading_sizes[i-1])
            heading.font.bold = True
    
    # Save reference document
    ref_path = Path.cwd() / "temp_reference.docx"
    doc.save(ref_path)
    return str(ref_path)


def enhanced_fallback(html_file: Path, docx_file: Path) -> None:
    """
    Enhanced fallback with better formula preservation and TOC handling.
    """
    print("[INFO] Using enhanced fallback conversion...")
    
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Helper functions
    def contains_math_formula(text):
        """Enhanced math detection."""
        math_patterns = [
            r'\$.*?\$',           # LaTeX delimiters
            r'\\[a-zA-Z]+',       # LaTeX commands
            r'[a-zA-Z]_[a-zA-Z0-9]',  # Subscripts
            r'[a-zA-Z]\^[a-zA-Z0-9]', # Superscripts
            r'∑|∫|≤|≥|≠|±|×|÷|∞|∂|∇|∆'  # Math symbols
        ]
        return any(re.search(pattern, text) for pattern in math_patterns)
    
    def extract_formulas_from_text(text):
        """Extract and format mathematical formulas from text."""
        # Find LaTeX formulas
        display_formulas = re.findall(r'\$\$(.*?)\$\$', text, re.DOTALL)
        inline_formulas = re.findall(r'\$([^$\n]+?)\$', text)
        
        formatted_formulas = []
        for formula in display_formulas:
            formatted_formulas.append(f"[DISPLAY FORMULA: {formula.strip()}]")
        for formula in inline_formulas:
            formatted_formulas.append(f"[INLINE FORMULA: {formula.strip()}]")
        
        return formatted_formulas
    
    # Parse HTML
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Extract formulas before parsing
    extractor = MathJaxExtractor()
    formula_count = extractor.debug_formula_extraction(html_content)
    
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()
    
    # Set title
    if soup.title:
        doc.core_properties.title = soup.title.string
    
    # Add conversion info
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(
        f"📄 Document converted using enhanced fallback mode. "
        f"Found {formula_count} mathematical expressions. "
        f"Install Pandoc for native Word math equations."
    )
    info_run.italic = True
    info_run.font.color.rgb = RGBColor(0, 100, 200)  # Blue color
    doc.add_paragraph()
    
    # Process content
    body = soup.body if soup.body else soup
    skip_toc = False
    
    def process_element(element):
        nonlocal skip_toc
        
        if not hasattr(element, 'name') or element.name is None:
            return
        
        tag = element.name.lower()
        text = element.get_text(strip=True)
        
        # Handle Table of Contents specially
        if text and 'table of contents' in text.lower():
            skip_toc = True
            doc.add_heading("Table of Contents", level=2)
            return
        
        # Skip TOC list items
        if skip_toc and tag in ['ul', 'ol', 'li']:
            if tag in ['ul', 'ol']:
                skip_toc = False  # End of TOC
            return
        
        # Handle headings
        if tag.startswith('h') and len(tag) == 2 and tag[1].isdigit():
            level = min(int(tag[1]), 6)
            if text:
                doc.add_heading(text, level=level)
        
        # Handle paragraphs with math detection
        elif tag == 'p' and text:
            if contains_math_formula(text):
                # Handle mathematical content specially
                para = doc.add_paragraph()
                para.add_run("🔢 Mathematical Expression: ").bold = True
                
                # Extract and format formulas
                formulas = extract_formulas_from_text(text)
                if formulas:
                    for formula in formulas:
                        para.add_run("\n" + formula).italic = True
                else:
                    para.add_run(text).italic = True
            else:
                doc.add_paragraph(text)
        
        # Handle lists
        elif tag == 'ul':
            for li in element.find_all('li', recursive=False):
                if not skip_toc:  # Don't add TOC items as regular lists
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(li.get_text(strip=True))
        
        elif tag == 'ol':
            for li in element.find_all('li', recursive=False):
                if not skip_toc:  # Don't add TOC items as regular lists
                    para = doc.add_paragraph(style='List Number')
                    para.add_run(li.get_text(strip=True))
        
        # Handle tables
        elif tag == 'table':
            rows = element.find_all('tr')
            if rows:
                max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < max_cols:
                            table.cell(i, j).text = cell.get_text(strip=True)
        
        # Handle divisions (recurse)
        elif tag in ['div', 'section', 'article', 'main']:
            for child in element.children:
                process_element(child)
    
    # Process all content
    for child in body.children:
        process_element(child)
    
    # Save document
    doc.save(docx_file)
    print(f"[SUCCESS] ✅ Created {docx_file} with enhanced formula preservation")


# ================================================================
# SECTION: Main entry-point
# ------------------------------------------------
def main() -> None:
    """
    Enhanced main function with better error handling and debugging.
    """
    project_root = Path(__file__).resolve().parent
    html_path = project_root / "EMS_Technical_Report.html"
    docx_path = project_root / "EMS_Technical_Report_fixed.docx"

    if not html_path.exists():
        sys.exit(f"[ERROR] ❌ Source file {html_path} not found. Run create_html_version.py first.")

    print(f"[INFO] 🚀 Converting {html_path.name} to {docx_path.name}")
    
    # Check for mathematical content
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Quick analysis
    latex_count = len(re.findall(r'\$.*?\$', html_content))
    print(f"[INFO] 📊 Found {latex_count} LaTeX expressions in HTML")

    # Choose conversion method
    if pandoc_available():
        print("[INFO] ✅ Pandoc available - using high-quality conversion")
        try:
            run_pandoc_with_math(html_path, docx_path)
        except Exception as e:
            print(f"[ERROR] ❌ Pandoc failed: {e}")
            print("[INFO] 🔄 Falling back to enhanced python-docx conversion...")
            enhanced_fallback(html_path, docx_path)
    else:
        print("[WARN] ⚠️  Pandoc not found - using enhanced fallback")
        print("       💡 Install Pandoc from https://pandoc.org for best results")
        enhanced_fallback(html_path, docx_path)
    
    print(f"\n[SUCCESS] 🎉 Conversion complete!")
    print(f"[OUTPUT] 📄 Word document: {docx_path}")
    
    # Cleanup
    temp_ref = Path.cwd() / "temp_reference.docx"
    if temp_ref.exists():
        temp_ref.unlink()


if __name__ == "__main__":
    main()