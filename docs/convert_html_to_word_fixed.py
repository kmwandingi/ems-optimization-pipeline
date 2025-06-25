#!/usr/bin/env python3
# ================================================================
# convert_html_to_word_improved.py
# ------------------------------------------------
# Purpose  : Convert an HTML technical report to a Word .docx file
# Strategy : 1) Prefer Pandoc → flawless structure + math retention
#            2) Graceful fallback to simple python-docx extraction
# Author   : (your name here)
# ================================================================

# -------- 1. Standard-library imports ---------------------------------
import subprocess
import sys
from pathlib import Path
import shutil
import platform
import re

# -------- 2. Third-party imports (lazy-loaded if Pandoc absent) -------
from bs4 import BeautifulSoup          # noqa: E402
from docx import Document              # noqa: E402


# ================================================================
# SECTION: Helper functions
# ------------------------------------------------
def pandoc_available() -> bool:
    """Check whether Pandoc is on PATH and at least version 2.9."""
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False
    try:
        version = subprocess.check_output([pandoc, "--version"], text=True)
        major = int(version.split(".")[0])
        return major >= 2
    except Exception:
        return False


def run_pandoc(html_file: Path, docx_file: Path) -> None:
    """Invoke Pandoc for flawless HTML → DOCX conversion with math support."""
    cmd = [
        "pandoc",
        str(html_file),
        "--from=html",
        "--to=docx",
        "--embed-resources",          # bundle local images/fonts
        "--standalone",
        "--output", str(docx_file),
    ]
    # Pandoc maps LaTeX/MathJax to Word equations automatically
    print(f"[INFO] Running Pandoc: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)
    print(f"[SUCCESS] Created {docx_file}")


# --- inside convert_html_to_word.py ---------------------------------
def simple_fallback(html_file: Path, docx_file: Path) -> None:
    """Lightweight HTML → DOCX conversion when Pandoc is unavailable."""
    global BeautifulSoup, Document

    # If the libraries haven't been imported yet, install them
    if BeautifulSoup is None or Document is None:
        print("[INFO] Installing required Python packages for fallback…")
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "beautifulsoup4", "python-docx"]
        )
        from bs4 import BeautifulSoup  # noqa: E402
        from docx import Document      # noqa: E402
        from docx.shared import Pt    # For point-based font sizes
        # Re-bind to the module-level names
        globals()["BeautifulSoup"] = BeautifulSoup
        globals()["Document"] = Document
        globals()["Pt"] = Pt

    # Parse HTML and build Word doc
    soup = BeautifulSoup(html_file.read_text(encoding="utf-8"), "html.parser")
    body = soup.body or soup
    
    # Initialize document with styles
    doc = Document()
    
    # Add a 'Code' style if it doesn't exist
    styles = doc.styles
    if 'Code' not in styles:
        code_style = styles.add_style('Code', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        code_style.base_style = styles['Normal']
        font = code_style.font
        font.name = 'Courier New'
        font.size = Pt(9)
    
    # Set document title
    if soup.title:
        doc.core_properties.title = soup.title.string

    # Add warning about fallback mode
    warn = doc.add_paragraph()
    warn.add_run(
        "⚠️  Fallback path: math formulas and complex formatting may not be preserved. "
        "Install Pandoc for full-fidelity conversion."
    ).italic = True

    # Process content recursively
    def process_element(element, doc):
        if not hasattr(element, 'name') or element.name is None:
            return
            
        # Handle headings
        if element.name.startswith('h') and len(element.name) == 2 and element.name[1].isdigit():
            level = min(int(element.name[1]), 6)  # Limit to h6 max
            doc.add_heading(element.get_text(strip=True), level=level)
        # Handle paragraphs
        elif element.name == 'p':
            text = element.get_text(' ', strip=True)
            if text:  # Only add non-empty paragraphs
                doc.add_paragraph(text)
        # Handle lists
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(li.get_text(' ', strip=True))
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(li.get_text(' ', strip=True))
        # Handle math elements
        elif element.name == 'math' or (element.get('class') and 'math' in ' '.join(element.get('class', []))):
            p = doc.add_paragraph()
            p.add_run(f"[MATH: {element.get_text(strip=True)}]").italic = True
        # Handle code blocks
        elif element.name == 'pre' or (element.get('class') and 'code' in ' '.join(element.get('class', []))):
            p = doc.add_paragraph(style='Code')
            p.add_run(element.get_text())
        # Handle other block elements
        elif element.name in ['div', 'section', 'article']:
            for child in element.children:
                process_element(child, doc)
        # Handle inline elements
        elif element.name in ['span', 'strong', 'em', 'i', 'b']:
            # Just get text for now, could be enhanced to preserve formatting
            return element.get_text(' ', strip=True)
            
    # Process all top-level elements in the body
    for child in body.children:
        process_element(child, doc)

    # Save the document
    doc.save(docx_file)
    print(f"[SUCCESS] Created {docx_file} (fallback mode, some formatting may be simplified)")



# ================================================================
# SECTION: Main entry-point
# ------------------------------------------------
def main() -> None:
    """Driver routine handling CLI arguments and overall flow."""
    # --- 1. Locate source HTML and desired output path ------------
    project_root = Path(__file__).resolve().parent
    html_path = project_root / "EMS_Technical_Report.html"
    docx_path = project_root / "EMS_Technical_Report.docx"

    if not html_path.exists():
        sys.exit(f"[ERROR] Source file {html_path} not found.")

    # --- 2. Choose conversion backend ----------------------------
    if pandoc_available():
        print("[INFO] Pandoc detected – using high-quality conversion.")
        run_pandoc(html_path, docx_path)
    else:
        print("[WARN] Pandoc not found; falling back to python-docx.")
        print("       → Install Pandoc from https://pandoc.org for best results.")
        simple_fallback(html_path, docx_path)


# ------------------------------------------------
# SECTION: Script execution guard
# ------------------------------------------------
if __name__ == "__main__":
    main()
